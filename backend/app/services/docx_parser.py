from typing import List, Dict, Any
from docx import Document


def extract_text_from_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Extract text content from a DOCX file with paragraph information
    
    Returns a list of dictionaries containing:
    - text: The paragraph text
    - paragraph_index: The index of the paragraph
    - style: The paragraph style name
    """
    doc = Document(file_path)
    paragraphs = []
    
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            paragraphs.append({
                "text": paragraph.text,
                "paragraph_index": idx,
                "style": paragraph.style.name if paragraph.style else "Normal"
            })
    
    return paragraphs


def get_document_stats(file_path: str) -> Dict[str, Any]:
    """Get document statistics"""
    doc = Document(file_path)
    
    total_paragraphs = len(doc.paragraphs)
    total_words = sum(len(p.text.split()) for p in doc.paragraphs)
    total_characters = sum(len(p.text) for p in doc.paragraphs)
    
    return {
        "paragraphs": total_paragraphs,
        "words": total_words,
        "characters": total_characters
    }

