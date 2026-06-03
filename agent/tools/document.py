"""
Document parsing tools for proofreading agents
"""
from typing import Any
from docx import Document
from docx.opc.exceptions import PackageNotFoundError


def get_document_sections(file_path: str) -> list[dict[str, Any]]:
    """
    Parse document into sections by heading styles.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        List of sections, each containing:
        - heading: The section heading text
        - heading_level: 1-9 for Heading 1-9, 0 for no heading
        - content: List of paragraph texts in this section
        - paragraph_indices: List of paragraph indices in the original document
        - start_index: First paragraph index of this section
        - end_index: Last paragraph index of this section
    """
    try:
        doc = Document(file_path)
    except (PackageNotFoundError, Exception) as e:
        return [{"error": f"Failed to open document: {str(e)}"}]
    
    sections = []
    current_section = {
        "heading": "",
        "heading_level": 0,
        "content": [],
        "paragraph_indices": [],
        "start_index": 0,
        "end_index": 0
    }
    
    for idx, paragraph in enumerate(doc.paragraphs):
        style_name = paragraph.style.name if paragraph.style else "Normal"
        text = paragraph.text.strip()
        
        # Check if this is a heading
        if style_name.startswith("Heading"):
            # Save previous section if it has content
            if current_section["content"] or current_section["heading"]:
                current_section["end_index"] = idx - 1 if idx > 0 else 0
                sections.append(current_section)
            
            # Extract heading level
            try:
                level = int(style_name.replace("Heading ", "").strip())
            except ValueError:
                level = 1
            
            # Start new section
            current_section = {
                "heading": text,
                "heading_level": level,
                "content": [],
                "paragraph_indices": [idx],
                "start_index": idx,
                "end_index": idx
            }
        elif text:
            # Add content to current section
            current_section["content"].append(text)
            current_section["paragraph_indices"].append(idx)
    
    # Don't forget the last section
    if current_section["content"] or current_section["heading"]:
        current_section["end_index"] = len(doc.paragraphs) - 1
        sections.append(current_section)
    
    return sections


def get_paragraph(file_path: str, index: int) -> dict[str, Any]:
    """
    Get specific paragraph with surrounding context.
    
    Args:
        file_path: Path to the DOCX file
        index: Paragraph index (0-based)
        
    Returns:
        Dictionary containing:
        - text: The paragraph text
        - index: The paragraph index
        - style: The paragraph style name
        - prev_text: Previous paragraph text (if exists)
        - next_text: Next paragraph text (if exists)
        - prev_index: Previous paragraph index
        - next_index: Next paragraph index
    """
    try:
        doc = Document(file_path)
    except (PackageNotFoundError, Exception) as e:
        return {"error": f"Failed to open document: {str(e)}"}
    
    paragraphs = doc.paragraphs
    
    if index < 0 or index >= len(paragraphs):
        return {"error": f"Paragraph index {index} out of range (0-{len(paragraphs)-1})"}
    
    result = {
        "text": paragraphs[index].text,
        "index": index,
        "style": paragraphs[index].style.name if paragraphs[index].style else "Normal",
        "prev_text": None,
        "next_text": None,
        "prev_index": None,
        "next_index": None
    }
    
    # Get previous non-empty paragraph
    for i in range(index - 1, -1, -1):
        if paragraphs[i].text.strip():
            result["prev_text"] = paragraphs[i].text
            result["prev_index"] = i
            break
    
    # Get next non-empty paragraph
    for i in range(index + 1, len(paragraphs)):
        if paragraphs[i].text.strip():
            result["next_text"] = paragraphs[i].text
            result["next_index"] = i
            break
    
    return result


def get_document_metadata(file_path: str) -> dict[str, Any]:
    """
    Extract document metadata.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Dictionary containing:
        - title: Document title from properties
        - author: Document author
        - created: Creation date
        - modified: Last modified date
        - word_count: Approximate word count
        - paragraph_count: Total paragraph count
        - section_count: Number of sections (based on headings)
    """
    try:
        doc = Document(file_path)
    except (PackageNotFoundError, Exception) as e:
        return {"error": f"Failed to open document: {str(e)}"}
    
    core_props = doc.core_properties
    
    # Count words and paragraphs
    word_count = 0
    paragraph_count = 0
    heading_count = 0
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraph_count += 1
            word_count += len(text.split())
            
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name.startswith("Heading"):
                heading_count += 1
    
    return {
        "title": core_props.title or "",
        "author": core_props.author or "",
        "created": str(core_props.created) if core_props.created else None,
        "modified": str(core_props.modified) if core_props.modified else None,
        "word_count": word_count,
        "paragraph_count": paragraph_count,
        "section_count": heading_count + 1  # +1 for content before first heading
    }


def get_full_text(file_path: str) -> dict[str, Any]:
    """
    Extract full document text with paragraph information.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Dictionary containing:
        - text: Full document text (paragraphs joined by newlines)
        - paragraphs: List of {index, text, style} for each paragraph
    """
    try:
        doc = Document(file_path)
    except (PackageNotFoundError, Exception) as e:
        return {"error": f"Failed to open document: {str(e)}"}
    
    paragraphs = []
    text_parts = []
    
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        if text.strip():
            paragraphs.append({
                "index": idx,
                "text": text,
                "style": paragraph.style.name if paragraph.style else "Normal"
            })
            text_parts.append(text)
    
    return {
        "text": "\n".join(text_parts),
        "paragraphs": paragraphs
    }




